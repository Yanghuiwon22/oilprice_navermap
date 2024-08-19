import io
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor


from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph

from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from PIL import Image


from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service

chrome_driver_path = "/usr/bin/chromedriver"
service = Service(chrome_driver_path)

def outo_screenshot_km(start_location, end_location, waypoints):
    op = Options()
    op.add_argument('headless')
    op.add_argument('window-size=1920x1080')

    op.add_argument('--no-sandbox')

    op.add_argument('--disable-dev-shm-usage')


    op.add_argument("disable-gpu")
    op.add_argument(
        "user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36 Edg/126.0.0.0")
    op.add_argument('lang=ko_KR')

    try:
        print('server side')
        browser = webdriver.Chrome(service=service, options=op)
    except:
        print('local side')
        browser = webdriver.Chrome(options=op)
    url = 'https://map.naver.com/p?c=15.00,0,0,0,dh'
    browser.get(url)

    browser.maximize_window()

    side_var_button = WebDriverWait(browser, 10).until(
        EC.element_to_be_clickable((By.XPATH, "(//button[contains(@class, 'btn_navbar')])[2]"))
    )
    side_var_button.click()

    car_button = WebDriverWait(browser, 10).until(
        EC.element_to_be_clickable((By.XPATH, "(//button[contains(@class, 'btn_search_tab')])[2]"))
    )
    car_button.click()


    if len(waypoints) != 0:
        for i in range(len(waypoints)):
            waypoints_search = browser.find_element(By.CSS_SELECTOR, '.search_btn_area button:nth-of-type(2)')
            waypoints_search.click()

        search = browser.find_elements(By.CLASS_NAME, "input_search")

        for idx, waypoint in enumerate(waypoints):
            search[idx + 1].send_keys(waypoint)
            time.sleep(2)
            search[idx + 1].send_keys(Keys.RETURN)

        search[-1].send_keys(f"{end_location}")
        time.sleep(1.5)
        search[-1].send_keys(Keys.RETURN)
    else:
        search = browser.find_elements(By.CLASS_NAME, "input_search")
        search[1].send_keys(f"{end_location}")
        time.sleep(1.5)
        search[1].send_keys(Keys.RETURN)

    search[0].send_keys(f"{start_location}")
    time.sleep(2)
    search[0].send_keys(Keys.RETURN)

    time.sleep(1.5)
    road_search = browser.find_element(By.XPATH, '//*[@id="section_content"]/div/div[1]/div[2]/button[3]')
    road_search.click()
    time.sleep(5)


    text = browser.page_source
    print(type(text))
    distance_between_locations = browser.find_element(By.XPATH, '//*[@id="section_content"]/div/div[2]/div/div[2]/ul/li[1]/div/div/div[2]/span')
    distance = distance_between_locations.text
    print(f"총 거리 : {distance}")

    time.sleep(1)

    map_img = browser.find_element(By.XPATH, '//*[@id="app-layout"]')
    map_img.screenshot('./output/naver_map.png')

    return distance

def get_docx(start_location, end_location, waypoints, distance, oil_date, oil_price, color):
    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run('[여비증빙]')

    if waypoints != []:
        waypoints_text = []
        for waypoint in waypoints:
            text = f'{waypoint} ->'
            waypoints_text.append(text)
        run_route = doc.add_paragraph()
        route_text = run_route.add_run(f'{start_location} -> {" ".join(waypoints_text)} {end_location}')
        route_text.bold = True
        route_text.font.color.rgb = RGBColor(color[0], color[1], color[2])

        para_distance = doc.add_paragraph(f'총 거리 : ')
        distance_text = para_distance.add_run(distance)
        distance_text.bold = True
        distance_text.font.color.rgb = RGBColor(color[0], color[1], color[2])
        #para_distance.add_run('km')

    else:
        run_route = doc.add_paragraph()
        route_text = run_route.add_run(f'{start_location} -> {end_location}')
        route_text.bold = True
        route_text.font.color.rgb = RGBColor(color[0], color[1], color[2])

        para_distance = doc.add_paragraph(f'총 거리 : ')
        distance_text = para_distance.add_run(distance)
        distance_text.bold = True
        distance_text.font.color.rgb = RGBColor(color[0], color[1], color[2])

    image_path = './output/naver_map.png'
    doc.add_picture(image_path, width=Inches(5.0))

    para_oilprice = doc.add_paragraph()
    oildate_text = para_oilprice.add_run(str(oil_date))   # ---------> oil_date
    oildate_text.bold = True
    oildate_text.font.color.rgb = RGBColor(color[0], color[1], color[2])
    para_oilprice.add_run('의 휘발유 가격 : ')
    oilprice_text = para_oilprice.add_run(oil_price) # ---------> oil_price
    oilprice_text.bold = True
    oilprice_text.font.color.rgb = RGBColor(color[0], color[1], color[2])
    para_oilprice.add_run('원')

    image_path_oil = './output/oil_price.png'
    doc.add_picture(image_path_oil, width=Inches(5.0))

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def get_pdf(start_location, end_location, waypoints, distance, oil_date, oil_price, color):
    # PDF 파일을 메모리 버퍼에 저장합니다.
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.setFillColorRGB(color[0], color[1], color[2]) # RGB 색상을 설정합니다.

    font_path = "malgun.ttf"
    pdfmetrics.registerFont(TTFont("맑은고딕", font_path))
    pdfmetrics.registerFont(TTFont("맑은고딕-Bold", "malgunbd.ttf"))  # 볼드체 폰트 등록

    pdf.setFont("맑은고딕", 14)
    pdf.setFillColorRGB(0, 0, 0)
    pdf.drawString(50, 810, "[여비증빙]")

    pdf.setFont("맑은고딕-Bold", 12)
    pdf.setFillColorRGB(color[0], color[1], color[2]) # RGB 색상을 설정합니다.

    x_position = 50
    y_position = 780

    if waypoints != []:
        waypoints_text = []
        for waypoint in waypoints:
            text = f'{waypoint} ->'
            waypoints_text.append(text)

        pdf.drawString(50, y_position, f'{start_location} -> {" ".join(waypoints_text)} {end_location}')
        y_position -= 12*2.4

        pdf.setFont("맑은고딕", 12)
        pdf.setFillColorRGB(0,0,0)  # 검정
        pdf.drawString(50, y_position, f'총 거리 : ', )

        pdf.setFont("맑은고딕-Bold", 12)
        pdf.setFillColorRGB(color[0], color[1], color[2])  # RGB 색상을 설정합니다.
        pdf.drawString(100, y_position, f'{distance}')
        y_position -= 12*1.2

    else:
        pdf.drawString(50, y_position, f'{start_location} -> {end_location}')
        y_position -= 12*2.4

        pdf.setFont("맑은고딕", 12)
        pdf.setFillColorRGB(0,0,0)  # 검정
        pdf.drawString(50, y_position, f'총 거리 : ', )

        pdf.setFont("맑은고딕-Bold", 12)
        pdf.setFillColorRGB(color[0], color[1], color[2])  # RGB 색상을 설정합니다.
        pdf.drawString(100, y_position, f'{distance}')
        y_position -= 12 * 1.2


    ratio = get_image_ratio('./output/naver_map.png')
    img_width = 500
    y_position -= img_width/ratio
    pdf.drawImage('./output/naver_map.png', 50, y_position, img_width, img_width/ratio)

    y_position -= 12 * 2.4
    pdf.setFillColorRGB(color[0], color[1], color[2])  # RGB 색상을 설정합니다.
    text_width = pdf.stringWidth(f'{oil_date}')
    pdf.drawString(50, y_position, f'{oil_date}')

    pdf.setFont("맑은고딕", 12)
    x_position += text_width
    pdf.setFillColorRGB(0, 0, 0)  # 검정
    pdf.drawString(x_position, y_position, '의 휘발유 가격 : ')

    text_width = pdf.stringWidth('의 휘발유 가격 : ')
    x_position += text_width
    pdf.setFont("맑은고딕-Bold", 12)
    pdf.setFillColorRGB(color[0], color[1], color[2])  # RGB 색상을 설정합니다.
    pdf.drawString(x_position, y_position, f'{oil_price}')

    pdf.setFont("맑은고딕", 12)
    text_width = pdf.stringWidth(f'{oil_price}')
    x_position += text_width
    pdf.setFillColorRGB(0, 0, 0)  # 검정
    pdf.drawString(x_position, y_position, '원')

    y_position -= 12 * 1.2
    ratio = get_image_ratio('./output/oil_price.png')
    img_width = 500
    y_position -= img_width/ratio
    pdf.drawImage('./output/oil_price.png', 50, y_position, img_width, img_width/ratio)

    pdf.save()

    # Save PDF to a file
    with open('./output/navermap_oilprice.pdf', 'wb') as f:
        f.write(buffer.getvalue())

    buffer.seek(0)
    return buffer


def get_image_ratio(image_path):
    # 이미지를 열어 너비와 높이를 가져옵니다.
    with Image.open(image_path) as img:
        width, height = img.size

    # 비율을 계산합니다.
    ratio = width / height
    return ratio

def main():
    get_pdf('전북대', '서울대', ['부산대', '대구대'], '100km', '2021-07-01', '2000')

if __name__ == "__main__":
    main()